# -*- coding: utf-8 -*-
"""Genera el documento Word de la práctica de Firestore (juguetería).

Sigue el esquema pedido: Título, Objetivo, Introducción, Materiales,
Desarrollo, Resultados, Análisis, Conclusión, Bibliografía.

Enfoque: los alumnos se CONECTAN a un Firestore que ya existe
(proyecto `nosqlpractica`). Hacen CRUD en la colección de su equipo
(equipoN); las queries se practican sobre `catalogo` (20 juguetes ya
cargados) y el tiempo real sobre `vivo` (colección compartida).
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0xE1, 0x1D, 0x48)   # rosa carrusel
BLUE = RGBColor(0x25, 0x63, 0xEB)
DARK = RGBColor(0x27, 0x27, 0x2A)
CODE_BG = "F2F2F5"

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = DARK


def shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def code(text):
    for line in text.strip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Inches(0.2)
        shade(p, CODE_BG)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def body(text, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    p.paragraph_format.space_after = Pt(8)
    return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            r = p.add_run(it[0]); r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def numbered(items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        if isinstance(it, tuple):
            r = p.add_run(it[0]); r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def h(text, level=1, color=ACCENT):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = color
    return p


# ============================================================
# PORTADA
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Práctica: Base de datos NoSQL con Firestore")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Juguetería virtual “Toys ForUs” — Esquema, CRUD, Queries y Tiempo real")
r.italic = True
r.font.size = Pt(13)
r.font.color.rgb = BLUE

doc.add_paragraph()

meta = doc.add_table(rows=6, cols=2)
meta.style = "Light List Accent 1"
campos = [
    ("Materia / Asignatura", "Bases de datos NoSQL"),
    ("Tema", "Firestore: modelado de documentos y operaciones"),
    ("Proyecto Firebase", "nosqlpractica (ya configurado por el profesor)"),
    ("Equipo / Colección CRUD", "equipo ____   (ej. equipo1, equipo2, …)"),
    ("Integrantes", "1) ______________  2) ______________  3) ______________"),
    ("Fecha de entrega", "____ / ____ / ________"),
]
for i, (k, v) in enumerate(campos):
    meta.rows[i].cells[0].text = k
    meta.rows[i].cells[1].text = v
    meta.rows[i].cells[0].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# ============================================================
# 1. TÍTULO
# ============================================================
h("1. Título", 1)
body("Conexión a una base de datos Firestore en la nube para operar una "
     "juguetería virtual: implementación de la capa de repositorio (CRUD, "
     "consultas y tiempo real) sobre colecciones de documentos.")

# ============================================================
# 2. OBJETIVO
# ============================================================
h("2. Objetivo", 1)
body("Que el alumno comprenda y aplique el modelo de datos NoSQL orientado a "
     "documentos usando Cloud Firestore. En esta práctica NO se crea una base "
     "de datos: el equipo se CONECTA a un Firestore que ya existe y completa "
     "el código que lo opera. Al terminar, el alumno será capaz de:")
bullets([
    ("Modelar un documento: ", "definir el esquema (schema) de un juguete y "
     "entender la denormalización."),
    ("Operar datos (CRUD): ", "crear, leer, actualizar y eliminar documentos "
     "en la colección de su equipo."),
    ("Consultar (Queries): ", "filtrar, ordenar y paginar con where, orderBy, "
     "limit y startAfter, incluyendo índices compuestos."),
    ("Sincronizar en tiempo real: ", "suscribirse a cambios con onSnapshot, el "
     "equivalente a un websocket en Firestore."),
])

# ============================================================
# 3. INTRODUCCIÓN
# ============================================================
h("3. Introducción", 1)
body("Las bases de datos NoSQL no usan tablas ni filas. Cloud Firestore, de "
     "Google, organiza la información en colecciones que contienen documentos, "
     "y cada documento es esencialmente un objeto JSON con campos. No hay un "
     "esquema rígido impuesto por el motor: la flexibilidad la controla el "
     "programador desde el código y la seguridad desde las Reglas de Firestore.")
body("En esta práctica trabajamos sobre una juguetería virtual llamada "
     "“Toys ForUs”, conectada al proyecto de Firebase nosqlpractica que el "
     "profesor ya dejó configurado en src/lib/firebase.ts. Cada juguete vive "
     "en un único documento autocontenido —diseño denormalizado—; esto evita "
     "los JOIN típicos de SQL: una sola lectura trae todo lo necesario para "
     "pintar una ficha de producto.")
body("Para no estorbarse entre equipos, la base tiene VARIAS colecciones y "
     "cada sección de la app usa una distinta:")
bullets([
    ("CRUD → colección de tu equipo (equipo1, equipo2, …): ", "tu zona de "
     "trabajo; ahí creas, editas y borras. Empieza vacía."),
    ("Queries → colección catalogo: ", "ya viene cargada con 20 juguetes "
     "variados. Solo la lees para practicar consultas."),
    ("Tiempo real → colección vivo: ", "ya viene cargada y es COMPARTIDA por "
     "toda la clase; cuando alguien cambia algo, todos lo ven al instante."),
])
body("La aplicación (Next.js + React) ya está hecha; el alumno solo programa "
     "la carpeta src/repos/, que es la única capa que habla con Firestore.")

# ============================================================
# 4. MATERIALES / HERRAMIENTAS
# ============================================================
h("4. Materiales / Herramientas", 1)
bullets([
    ("Lenguaje: ", "TypeScript (modo strict) y TSX/React."),
    ("Framework: ", "Next.js 16 (App Router) con React 19."),
    ("Base de datos: ", "Google Cloud Firestore (SDK Web de Firebase), proyecto "
     "nosqlpractica ya configurado."),
    ("Estilos: ", "Tailwind CSS v4."),
    ("Entorno de ejecución: ", "Node.js 18+ y npm."),
    ("IDE recomendado: ", "Visual Studio Code (con la extensión ESLint)."),
    ("Navegador: ", "Chrome / Edge / Firefox para probar la app."),
    ("Conexión a internet: ", "para leer y escribir en Firestore en la nube."),
])

# ============================================================
# 5. DESARROLLO
# ============================================================
h("5. Desarrollo", 1)
body("El procedimiento se divide en una preparación y cinco actividades. Cada "
     "función del repositorio que falta por implementar lanza un error visible "
     "“⛔ ACTIVIDAD N…” en la pantalla, hasta que el alumno la completa. "
     "Importante: cada función del repositorio recibe como PRIMER parámetro el "
     "nombre de la colección sobre la que trabaja.")

h("Preparación del entorno", 2, BLUE)
numbered([
    "Descargar el proyecto y abrir la carpeta en VS Code.",
    "Instalar dependencias y arrancar el servidor de desarrollo:",
])
code("npm install\nnpm run dev      # abre http://localhost:3000")
body("No hay que crear ningún proyecto de Firebase ni capturar claves: ya "
     "están en src/lib/firebase.ts (son identificadores públicos, no secretos).",
     italic=True)

# --- Actividad 0 ---
h("Actividad 0 · Personalizar la tienda y elegir tu colección", 2, BLUE)
body("Abre src/config/store.config.ts. Cambia el nombre de la juguetería, el "
     "eslogan y los integrantes. Lo MÁS importante: pon en collections.crud la "
     "colección de TU equipo (equipo1, equipo2, …). Así tu CRUD no se mezcla "
     "con el de otros equipos.")
code('''export const storeConfig: StoreConfig = {
  storeName: "Tu Juguetería",
  tagline: "Tu eslogan mágico ✨",
  team: [{ name: "Tu Nombre", role: "Matrícula" }],
  categories: ["Peluches", "Vehículos", "Educativos", /* ... */],
  collections: {
    crud: "equipo1",      // 👈 CÁMBIALA por la de tu equipo
    catalog: "catalogo",  // 20 juguetes ya cargados (Queries)
    live: "vivo",         // colección compartida (Tiempo real)
  },
  theme: { primary: "#e11d48", secondary: "#2563eb", accent: "#f59e0b" },
};''')

# --- Actividad 1 ---
h("Actividad 1 · El esquema (schema)", 2, BLUE)
body("Archivo: src/repos/product.schema.ts. Define la forma del documento que "
     "vamos a LEER. Un Product es el documento guardado; un ProductInput es lo "
     "que captura el formulario (sin id ni metadatos del servidor).")
code('''import { Timestamp } from "firebase/firestore";

export interface Product {
  id: string;            // ID del documento (no se guarda dentro)
  name: string;
  description: string;
  category: string;
  price: number;
  stock: number;
  imageUrl?: string;     // opcionales
  ageRange?: string;
  brand?: string;
  createdAt?: Timestamp; // los pone el servidor
  updatedAt?: Timestamp;
}

// Lo que escribe el usuario = Product sin id ni metadatos.
export type ProductInput = Omit<Product, "id" | "createdAt" | "updatedAt">;''')

# --- Actividad 2 ---
h("Actividad 2 · CRUD (sobre la colección de tu equipo)", 2, BLUE)
body("Archivo: src/repos/products.repo.ts. Implementa las cuatro operaciones. "
     "Fíjate que la colección llega como parámetro (collectionName). Se prueba "
     "en la sección “El Mostrador”, que ya pasa la colección de tu equipo.")
code('''import { db } from "@/lib/firebase";
// helper ya incluido:
function colRef(collectionName: string) {
  return collection(db, collectionName);
}

// CREATE
export async function createProduct(collectionName: string, input: ProductInput) {
  return addDoc(colRef(collectionName), {
    ...input,
    createdAt: serverTimestamp(),
    updatedAt: serverTimestamp(),
  });
}

// UPDATE
export async function updateProduct(collectionName: string, id: string, input: ProductInput) {
  return updateDoc(doc(db, collectionName, id), {
    ...input,
    updatedAt: serverTimestamp(),
  });
}

// DELETE
export async function deleteProduct(collectionName: string, id: string) {
  return deleteDoc(doc(db, collectionName, id));
}

// READ (lectura puntual, sin tiempo real)
export async function fetchAllProducts(collectionName: string): Promise<Product[]> {
  const snap = await getDocs(query(colRef(collectionName), orderBy("createdAt", "desc")));
  return snap.docs.map(toProduct);
}''')

# --- Actividad 3 ---
h("Actividad 3 · Queries (sobre la colección catalogo)", 2, BLUE)
body("Misma archivo. Implementa fetchProductsPage combinando where, orderBy, "
     "limit y startAfter sobre la colección catalogo (ya tiene 20 juguetes). "
     "Truco: pedimos un documento de más (limit pageSize + 1) para saber si "
     "hay página siguiente.")
code('''export async function fetchProductsPage(
  collectionName: string,
  params: PageParams
): Promise<PageResult> {
  const { category, sortDir, pageSize, cursor } = params;

  const constraints: QueryConstraint[] = [];
  if (category) constraints.push(where("category", "==", category));
  constraints.push(orderBy("price", sortDir));
  if (cursor) constraints.push(startAfter(cursor));
  constraints.push(limit(pageSize + 1)); // uno de más

  const snap = await getDocs(query(colRef(collectionName), ...constraints));
  const docs = snap.docs;
  const hasMore = docs.length > pageSize;
  const pageDocs = hasMore ? docs.slice(0, pageSize) : docs;

  return {
    products: pageDocs.map(toProduct),
    lastDoc: pageDocs.length ? pageDocs[pageDocs.length - 1] : null,
    hasMore,
  };
}''')
body("La primera vez que filtres por categoría Y ordenes por precio, Firestore "
     "pedirá un ÍNDICE COMPUESTO y mostrará un enlace en el error: ábrelo, "
     "presiona “Crear índice” y espera unos segundos. Es parte del aprendizaje.",
     italic=True)

# tabla de los 20 juguetes de catalogo
body("Los 20 juguetes ya cargados en la colección catalogo (para que sepas qué "
     "estás consultando):", bold=True)
CATALOGO = [
    ("Oso Abrazos de peluche", "Peluches", 349),
    ("Conejo Saltarín de peluche", "Peluches", 279),
    ("Dragón Dormilón XL", "Peluches", 599),
    ("Bloques Arcoíris 100 pzas", "Bloques y construcción", 459),
    ("Set Castillo Medieval 250 pzas", "Bloques y construcción", 899),
    ("Cubos Mágicos Apilables", "Bloques y construcción", 199),
    ("Lotería Mexicana Clásica", "Juegos de mesa", 149),
    ("Serpientes y Escaleras Deluxe", "Juegos de mesa", 189),
    ("Ajedrez de Madera", "Juegos de mesa", 399),
    ("Memorama de Animales", "Juegos de mesa", 129),
    ("Carrito Bólido R/C", "Vehículos", 549),
    ("Tren Expreso de Madera", "Vehículos", 429),
    ("Avión Planeador de Espuma", "Vehículos", 99),
    ("Camión de Bomberos con Luces", "Vehículos", 329),
    ("Súper Robot Articulado", "Figuras de acción", 379),
    ("Dino T-Rex Rugiente", "Figuras de acción", 299),
    ("Set Cocina Educativa", "Educativos", 689),
    ("Microscopio Junior", "Educativos", 749),
    ("Pizarrón Mágico ABC", "Educativos", 219),
    ("Cometa Dragón Gigante", "Aire libre", 179),
]
tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
for i, t in enumerate(["Juguete", "Categoría", "Precio (MXN)"]):
    hdr[i].text = t
    hdr[i].paragraphs[0].runs[0].bold = True
for name, cat, price in CATALOGO:
    row = tbl.add_row().cells
    row[0].text = name
    row[1].text = cat
    row[2].text = f"${price}"
doc.add_paragraph()

# --- Actividad 4 ---
h("Actividad 4 · Tiempo real (sobre la colección vivo)", 2, BLUE)
body("Misma archivo. Implementa subscribeProducts: en lugar de leer una vez, "
     "se SUSCRIBE a la colección vivo y recibe la lista cada vez que algo "
     "cambia, en cualquier dispositivo. Devuelve la función para cancelar la "
     "suscripción. Como vivo es compartida, verás los cambios de otros equipos "
     "en vivo. Se prueba en “La Vitrina en Vivo”.")
code('''export function subscribeProducts(
  collectionName: string,
  onChange: (products: Product[]) => void,
  onError?: (e: Error) => void
): Unsubscribe {
  return onSnapshot(
    query(colRef(collectionName), orderBy("createdAt", "desc")),
    (snap) => onChange(snap.docs.map(toProduct)),
    (err) => onError?.(err)
  );
}''')

# ============================================================
# 6. RESULTADOS
# ============================================================
h("6. Resultados", 1)
body("Al completar las actividades, la aplicación queda totalmente funcional. "
     "Resultados esperados por sección:")
bullets([
    ("El Mostrador (CRUD): ", "el formulario agrega juguetes a la colección de "
     "tu equipo, la lista los muestra, y Editar/Eliminar funcionan. Empieza "
     "vacía y se llena con lo que tú creas."),
    ("El Catálogo (Queries): ", "se ven los 20 juguetes de catalogo; los "
     "selectores filtran por categoría y ordenan por precio; los botones "
     "Anterior/Siguiente paginan de 6 en 6 (4 páginas)."),
    ("La Vitrina en Vivo (tiempo real): ", "el indicador muestra “EN VIVO”; al "
     "abrir la página en dos pestañas (o con otro equipo) y cambiar el stock, "
     "ambas se actualizan al instante y la tarjeta modificada parpadea."),
])
body("Espacio para capturas de pantalla (pega aquí tus evidencias):", bold=True)
for etiqueta in ["[ Captura 1 — El Mostrador (CRUD en la colección de tu equipo) ]",
                 "[ Captura 2 — El Catálogo (filtro/orden/paginación sobre catalogo) ]",
                 "[ Captura 3 — La Vitrina en Vivo (dos pestañas sincronizadas) ]"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade(p, "FAFAFA")
    r = p.add_run("\n" + etiqueta + "\n\n")
    r.italic = True
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ============================================================
# 7. ANÁLISIS
# ============================================================
h("7. Análisis", 1)
body("¿Cómo funciona cada pieza?")
bullets([
    ("Colecciones por equipo: ", "usar equipo1, equipo2, … aísla el trabajo de "
     "cada equipo en su propia “carpeta” de documentos; nadie pisa los datos "
     "de otro. Cambiar de colección es solo cambiar un string en la config."),
    ("Documento autocontenido: ", "todo el juguete vive en un solo documento. "
     "No hay JOINs; una lectura basta. A cambio, si un dato se repite en muchos "
     "documentos, hay que actualizarlo en cada uno (denormalización)."),
    ("Capa de repositorio: ", "los componentes nunca llaman a Firestore "
     "directo; solo usan funciones de src/repos/ a las que les pasan la "
     "colección. Si cambiáramos de base de datos, solo se tocaría esa carpeta."),
    ("serverTimestamp(): ", "la fecha la pone el servidor de Google, no el "
     "reloj (posiblemente mal ajustado) del cliente."),
    ("Paginación con cursor: ", "startAfter(últimoDoc) es más eficiente que "
     "saltar con offset, porque Firestore continúa justo donde quedó."),
    ("Índice compuesto: ", "filtrar y ordenar por campos distintos obliga a "
     "Firestore a tener un índice que combine ambos; por eso lo pide la "
     "primera vez."),
    ("onSnapshot vs getDocs: ", "getDocs lee una vez; onSnapshot mantiene una "
     "conexión y reacciona a cada cambio. Es el modelo de tiempo real "
     "(equivalente a un websocket); por eso la colección vivo se sincroniza "
     "sola entre todos los dispositivos de la clase."),
    ("Claves públicas + Reglas: ", "las claves del SDK Web no son secretas; la "
     "seguridad real se define en las Reglas de Firestore, no ocultando "
     "identificadores."),
])

# ============================================================
# 8. CONCLUSIÓN
# ============================================================
h("8. Conclusión", 1)
body("Con esta práctica se comprende el cambio de mentalidad de SQL a NoSQL: "
     "en lugar de normalizar en muchas tablas relacionadas, se modela un "
     "documento autocontenido pensado para cómo se va a leer, y la información "
     "se organiza en colecciones (aquí, una por equipo). Se implementó la capa "
     "de acceso completa —CRUD, consultas con filtro/orden/paginación e "
     "índices, y sincronización en tiempo real— separada de la interfaz, lo "
     "que demuestra el valor de una arquitectura por capas (repositorio). "
     "Finalmente, onSnapshot sobre la colección compartida deja claro por qué "
     "Firestore es tan usado en aplicaciones colaborativas: el tiempo real se "
     "obtiene casi gratis.")
body("Reflexión del equipo (qué fue lo más difícil y qué aprendieron):",
     bold=True)
for _ in range(3):
    p = doc.add_paragraph("_" * 95)
    p.paragraph_format.space_after = Pt(10)

# ============================================================
# 9. BIBLIOGRAFÍA
# ============================================================
h("9. Bibliografía", 1)
refs = [
    "Google. (2024). Cloud Firestore — Documentación oficial. "
    "https://firebase.google.com/docs/firestore",
    "Google. (2024). Agregar datos a Cloud Firestore. "
    "https://firebase.google.com/docs/firestore/manage-data/add-data",
    "Google. (2024). Realizar consultas a tus datos en Cloud Firestore. "
    "https://firebase.google.com/docs/firestore/query-data/queries",
    "Google. (2024). Obtener actualizaciones en tiempo real con Cloud "
    "Firestore. https://firebase.google.com/docs/firestore/query-data/listen",
    "Google. (2024). Administrar índices en Cloud Firestore. "
    "https://firebase.google.com/docs/firestore/query-data/indexing",
    "Vercel. (2024). Next.js Documentation — App Router. "
    "https://nextjs.org/docs",
    "Meta. (2024). React Documentation. https://react.dev",
]
for r in refs:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(r)

import os
out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                   "PRACTICA_Firestore_Jugueteria.docx")
doc.save(out)
print("OK ->", out)
